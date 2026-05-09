from __future__ import annotations

import argparse
import sys
from pathlib import Path

if __package__ in {None, ""}:
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from src.styles import (
    LIGHT_GREY,
    VERY_LIGHT_GREY,
    add_horizontal_rule,
    add_page_number,
    add_toc,
    PRIMARY_BLUE,
    set_cell_shading,
    set_document_defaults,
    set_table_borders,
)
from src.utils import (
    build_index_sections,
    build_validation_issues,
    generate_validation_report,
    load_kpi_dataset,
    resolve_workbook_source,
)

SUMMARY_FIELDS = [
    ("KPI Code", "KPI_Code"),
    ("KPI Name TH", "KPI_Name_TH"),
    ("KPI Name EN", "KPI_Name_EN"),
    ("KPI Dimension", "KPI_Dimension"),
    ("Direction", "Direction"),
    ("Unit", "Unit"),
    ("Frequency", "Frequency"),
    ("Owner Department", "Owner_Department"),
    ("Data Source", "Data_Source"),
    ("Service Area", "Service_Area"),
    ("Clinical Area", "Clinical_Area"),
    ("Related KPI", "Related_KPI_Code"),
]

DETAIL_FIELDS = [
    ("ชื่อ KPI", "KPI_Name_TH", "KPI_Name_EN"),
    ("ชื่อย่อ", "KPI_Short_Name", None),
    ("วัตถุประสงค์", "Objective_TH", "Objective_EN"),
    ("คำจำกัดความ", "Definition_TH", "Definition_EN"),
    ("สูตรการคำนวณ", "Formula", None),
    ("คำอธิบายตัวตั้ง", "Numerator_Definition", None),
    ("คำอธิบายตัวหาร", "Denominator_Definition", None),
    ("เกณฑ์การแปลผล", "Interpretation_TH", "Interpretation_EN"),
    ("ค่าเป้าหมาย", "Benchmark", None),
    ("คำค้นภาษาไทย", "Search_Keywords_TH", None),
    ("คำค้นภาษาอังกฤษ", "Search_Keywords_EN", None),
    ("คำพ้องไทย", "Synonym_TH", None),
    ("คำพ้องอังกฤษ", "Synonym_EN", None),
    ("Dashboard URL", "Dashboard_URL", None),
    ("Confluence URL", "Confluence_URL", None),
    ("เอกสารอ้างอิง", "Reference_Text_All", None),
    ("หมายเหตุ", "Note", None),
]

LOGIC_FIELDS = [
    ("Unit_of_Analysis", "Unit_of_Analysis"),
    ("Population_Logic", "Population_Logic"),
    ("Numerator_Logic", "Numerator_Logic"),
    ("Denominator_Logic", "Denominator_Logic"),
    ("Inclusion_Criteria", "Inclusion_Criteria"),
    ("Exclusion_Criteria", "Exclusion_Criteria"),
    ("Time_Window", "Time_Window"),
    ("Date_Field", "Date_Field"),
    ("Data_Source_Table", "Data_Source_Table"),
    ("Join_Key", "Join_Key"),
    ("SQL_Where_Clause", "SQL_Where_Clause"),
    ("Data_Quality_Check", "Data_Quality_Check"),
    ("Known_Limitation", "Known_Limitation"),
]

VERSION_FIELDS = [
    ("Version", "Version"),
    ("Effective_Date", "Effective_Date"),
    ("Last_Update_Date", "Last_Update_Date"),
    ("Revision_Reason", "Revision_Reason"),
    ("Changed_By", "Changed_By"),
    ("Approved_By", "Approved_By"),
]

DETAIL_TABLE_FIELDS = [
    ("หมวดตัวชี้วัด", "KPI_Category", None),
    ("ประเภทตัวชี้วัด", "KPI_Type", None),
    ("รหัสตัวชี้วัด", "KPI_Code", None),
    ("ชื่อตัวชี้วัด (ภาษาไทย)", "KPI_Name_TH", None),
    ("ชื่อตัวชี้วัด (ภาษาอังกฤษ)", "KPI_Name_EN", None),
    ("นิยาม คำอธิบาย ความหมายของตัวชี้วัด", "Definition_TH", "Definition_EN"),
    ("วัตถุประสงค์ของตัวชี้วัด", "Objective_TH", "Objective_EN"),
    ("สูตรในการคำนวณ", "Formula", None),
]

OPTIONAL_DETAIL_FIELDS = [
    ("ความถี่ในการจัดทำข้อมูล", "Frequency", None),
    ("Benchmark (แหล่งอ้างอิง/ปี)*", "Benchmark", None),
    ("วิธีการแปลผล", "Interpretation_TH", "Interpretation_EN"),
    ("ทีม/Reference", "Reference_Text_All", None),
    ("วัน เดือน ปี ที่เริ่มใช้", "Effective_Date", None),
    ("วัน เดือน ปี ที่ปรับปรุงครั้งล่าสุด", "Last_Update_Date", None),
    ("เหตุผลของการปรับปรุง", "Revision_Reason", None),
    ("หมายเหตุ", "Note", None),
]


def _ensure_heading_styles(document: Document, font_th: str, font_en: str) -> None:
    for style_name, size in (("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)):
        style = document.styles[style_name]
        style.font.name = font_en
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = PRIMARY_BLUE
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_th)

    if "KPI English" not in document.styles:
        style = document.styles.add_style("KPI English", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font_en
        style.font.size = Pt(10)
        style.font.italic = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_th)


def _set_run_fonts(run, font_th: str, font_en: str, size: int | None = None, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_en
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_th)


def _set_paragraph_font(paragraph, font_th: str, font_en: str, size: int = 11, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        _set_run_fonts(run, font_th, font_en, size=size, bold=bold, italic=italic)


def _compose_multilingual_value(record: dict[str, str], thai_key: str, english_key: str | None) -> str:
    thai_value = record.get(thai_key, "")
    english_value = record.get(english_key, "") if english_key else ""
    if thai_value and english_value:
        return f"{thai_value}\n{english_value}"
    return thai_value or english_value


def _add_cover_page(document: Document, config: dict[str, str]) -> None:
    logo_path = Path(config.get("Logo_Path", ""))
    if logo_path.exists():
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(logo_path), width=Cm(3))

    for text, size, bold in (
        (config.get("Document_Title", "KPI Dictionary"), 24, True),
        (config.get("Document_Subtitle", ""), 18, False),
        ("", 12, False),
        (config.get("Organization_Name", ""), 16, True),
        (f"Version {config.get('Version', '1.0')}", 12, False),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            run = paragraph.add_run(text)
            _set_run_fonts(run, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), size=size, bold=bold)


def _add_toc_page(document: Document, config: dict[str, str]) -> None:
    document.add_page_break()
    heading = document.add_paragraph("สารบัญ", style="Heading 1")
    add_horizontal_rule(heading)
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    note = document.add_paragraph("เปิดเอกสารใน Microsoft Word แล้วอัปเดต Table of Contents เพื่อแสดงเลขหน้าล่าสุด")
    _set_paragraph_font(note, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), size=10)


def _format_table_fonts(table, font_th: str, font_en: str, bold_first_column: bool = False) -> None:
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_fonts(run, font_th, font_en, size=11, bold=bold_first_column and index == 0)


def _add_index_section(document: Document, config: dict[str, str], title: str, rows: list[dict[str, str]]) -> None:
    heading = document.add_paragraph(title, style="Heading 2")
    add_horizontal_rule(heading, color="B7B7B7")

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_borders(table)
    for cell, label in zip(table.rows[0].cells, ["รหัส", "ชื่อตัวชี้วัด", "หน้า"], strict=False):
        cell.text = label
        set_cell_shading(cell, LIGHT_GREY)
        cell.paragraphs[0].runs[0].bold = True

    last_value = None
    for row_data in rows:
        if row_data["Index_Value"] != last_value:
            row = table.add_row().cells
            row[0].merge(row[1]).merge(row[2])
            row[0].text = row_data["Index_Value"]
            set_cell_shading(row[0], VERY_LIGHT_GREY)
            row[0].paragraphs[0].runs[0].bold = True
            last_value = row_data["Index_Value"]

        row = table.add_row().cells
        row[0].text = row_data.get("KPI_Code", "")
        row[1].paragraphs[0].text = row_data.get("KPI_Name_TH", "")
        en = row[1].add_paragraph(row_data.get("KPI_Name_EN", ""))
        en.style = "KPI English"
        row[2].text = row_data.get("Page_No", "")

    _format_table_fonts(table, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"))


def _add_all_index_pages(document: Document, config: dict[str, str], records: list[dict[str, str]]) -> None:
    document.add_page_break()
    heading = document.add_paragraph("Search Indexes", style="Heading 1")
    add_horizontal_rule(heading)
    for title, rows in build_index_sections(records).items():
        _add_index_section(document, config, title, rows)


def _add_tag_line(document: Document, config: dict[str, str], record: dict[str, str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tags = [
        record.get("KPI_Dimension", ""),
        record.get("Service_Area", ""),
        record.get("Clinical_Area", ""),
        record.get("Frequency", ""),
    ]
    for tag in [item for item in tags if item]:
        run = paragraph.add_run(f" {tag} ")
        _set_run_fonts(run, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), size=10, bold=True)
        run.font.color.rgb = PRIMARY_BLUE


def _add_summary_table(document: Document, config: dict[str, str], record: dict[str, str]) -> None:
    heading = document.add_paragraph("KPI Quick Summary", style="Heading 2")
    add_horizontal_rule(heading, color="D9D9D9")

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    for label, field in SUMMARY_FIELDS:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = record.get(field, "")
        set_cell_shading(row[0], VERY_LIGHT_GREY)
        if label == "KPI Name EN":
            row[1].paragraphs[0].style = "KPI English"
    _format_table_fonts(table, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), bold_first_column=True)


def _add_detail_table(document: Document, config: dict[str, str], record: dict[str, str]) -> None:
    heading = document.add_paragraph("KPI Dictionary Detail", style="Heading 2")
    add_horizontal_rule(heading, color="D9D9D9")

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    for label, thai_key, english_key in DETAIL_FIELDS:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = _compose_multilingual_value(record, thai_key, english_key)
        if label in {"สูตรการคำนวณ", "คำอธิบายตัวตั้ง", "คำอธิบายตัวหาร"}:
            set_cell_shading(row[0], LIGHT_GREY)
    _format_table_fonts(table, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), bold_first_column=True)


def _add_logic_and_version_table(document: Document, config: dict[str, str], record: dict[str, str]) -> None:
    heading = document.add_paragraph("Data Logic for Implementation", style="Heading 2")
    add_horizontal_rule(heading, color="D9D9D9")

    version_heading = document.add_paragraph("Version / Revision History", style="Heading 2")
    add_horizontal_rule(version_heading, color="D9D9D9")

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)

    section_row = table.add_row().cells
    section_row[0].merge(section_row[1])
    section_row[0].text = "Data Logic for Implementation"
    set_cell_shading(section_row[0], LIGHT_GREY)

    for label, field in LOGIC_FIELDS:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = record.get(field, "")

    version_row = table.add_row().cells
    version_row[0].merge(version_row[1])
    version_row[0].text = "Version / Revision History"
    set_cell_shading(version_row[0], LIGHT_GREY)

    for label, field in VERSION_FIELDS:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = record.get(field, "")

    _format_table_fonts(table, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), bold_first_column=True)


def _add_reference_style_detail_table(document: Document, config: dict[str, str], record: dict[str, str]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)

    def add_row(label: str, value: str, section: bool = False) -> None:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        if section:
            set_cell_shading(row[0], LIGHT_GREY)
            set_cell_shading(row[1], LIGHT_GREY)
        else:
            set_cell_shading(row[0], VERY_LIGHT_GREY)

    for label, thai_key, english_key in DETAIL_TABLE_FIELDS:
        add_row(label, _compose_multilingual_value(record, thai_key, english_key))

    add_row("ข้อมูลที่ต้องการ", "ข้อมูลที่ต้องการ", section=True)
    add_row("ตัวตั้ง", record.get("Numerator_Definition", ""))
    add_row("ตัวหาร", record.get("Denominator_Definition", ""))

    add_row("รหัสโรค/หัตถการที่เกี่ยวข้อง", "รหัสโรค/หัตถการที่เกี่ยวข้อง", section=True)
    add_row("ตัวตั้ง", record.get("Numerator_CodeSet", ""))
    add_row("ตัวหาร", record.get("Denominator_CodeSet", ""))

    for label, thai_key, english_key in OPTIONAL_DETAIL_FIELDS:
        add_row(label, _compose_multilingual_value(record, thai_key, english_key))

    _format_table_fonts(table, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), bold_first_column=True)


def _add_kpi_detail_page(document: Document, config: dict[str, str], record: dict[str, str]) -> None:
    document.add_page_break()

    code_line = document.add_paragraph()
    code_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = code_line.add_run(record.get("KPI_Code", ""))
    _set_run_fonts(run, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), size=20, bold=True)

    category_line = document.add_paragraph(f"หมวดตัวชี้วัด: {record.get('KPI_Category', '')}")
    _set_paragraph_font(category_line, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), size=11, bold=True)

    type_line = document.add_paragraph(f"ประเภทตัวชี้วัด: {record.get('KPI_Type', '')}")
    add_horizontal_rule(type_line)
    _set_paragraph_font(type_line, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), size=11)

    _add_reference_style_detail_table(document, config, record)


def _configure_footer(document: Document, config: dict[str, str]) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_run = footer.add_run(f"Version {config.get('Version', '1.0')} | Page ")
        _set_run_fonts(version_run, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"), size=10)
        add_page_number(footer)


def try_export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert
    except Exception:
        return False

    try:
        convert(str(docx_path), str(pdf_path))
    except Exception:
        return False
    return True


def generate_kpi_dictionary(
    workbook_path: Path | str,
    output_path: Path,
    export_pdf: bool = False,
    validation_report_path: Path | None = None,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    resolved_workbook_path = resolve_workbook_source(workbook_path, Path("temp"))

    dataset = load_kpi_dataset(resolved_workbook_path)
    issues = build_validation_issues(dataset)
    if validation_report_path is not None:
        generate_validation_report(validation_report_path, issues)

    config = dataset.config
    document = Document()
    set_document_defaults(document, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"))
    _ensure_heading_styles(document, config.get("Font_TH", "TH Sarabun New"), config.get("Font_EN", "Arial"))

    _add_cover_page(document, config)
    _add_toc_page(document, config)
    _add_all_index_pages(document, config, dataset.records)
    for record in dataset.records:
        _add_kpi_detail_page(document, config, record)
    _configure_footer(document, config)

    document.save(output_path)
    if export_pdf:
        try_export_pdf(output_path, output_path.with_suffix(".pdf"))
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate KPI Dictionary Word document from an Excel workbook.")
    parser.add_argument("workbook", nargs="?", default=None, help="Path to a local .xlsx workbook.")
    parser.add_argument(
        "--google-sheet",
        type=str,
        default=None,
        help="Public Google Sheet URL or sheet ID. The workbook will be exported to temp/ as .xlsx before processing.",
    )
    parser.add_argument("--output", type=Path, default=Path("output/KPI_Dictionary.docx"), help="Target .docx file path.")
    parser.add_argument("--pdf", action="store_true", help="Attempt to export a PDF alongside the .docx output.")
    parser.add_argument(
        "--validation-report",
        type=Path,
        default=Path("output/validation_report.xlsx"),
        help="Target validation report .xlsx file path.",
    )
    args = parser.parse_args()

    workbook_source = args.google_sheet or args.workbook
    if not workbook_source:
        parser.error("Provide either a local workbook path or --google-sheet.")

    output_path = generate_kpi_dictionary(
        workbook_source,
        args.output,
        export_pdf=args.pdf,
        validation_report_path=args.validation_report,
    )
    print(f"Generated: {output_path}")
    print(f"Validation report: {args.validation_report}")
    if args.pdf:
        print(f"PDF export attempted: {output_path.with_suffix('.pdf')}")


if __name__ == "__main__":
    main()
