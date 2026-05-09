from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from src.generate_document import generate_kpi_dictionary
from src.generate_template import SHEET_NAMES
from src.generate_template import create_template_workbook


class DocumentGenerationTests(unittest.TestCase):
    @staticmethod
    def _find_tables_with_label(document: Document, label: str) -> list:
        matched = []
        for table in document.tables:
            labels = [table.cell(row_index, 0).text.strip() for row_index in range(len(table.rows))]
            if label in labels:
                matched.append(table)
        return matched

    def test_generate_kpi_dictionary_creates_search_indexes_and_reference_style_detail_pages(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            template_path = Path(tmpdir) / "template.xlsx"
            output_path = Path(tmpdir) / "KPI_Dictionary.docx"
            report_path = Path(tmpdir) / "validation_report.xlsx"

            create_template_workbook(template_path)
            generate_kpi_dictionary(template_path, output_path, validation_report_path=report_path)

            self.assertTrue(output_path.exists())
            self.assertTrue(report_path.exists())

            document = Document(output_path)
            paragraph_text = "\n".join(p.text for p in document.paragraphs)

            self.assertIn("KPI Dictionary", paragraph_text)
            self.assertIn("Index by KPI Code", paragraph_text)
            self.assertIn("Index by KPI Category", paragraph_text)
            self.assertIn("Data Completeness Summary", paragraph_text)
            self.assertIn("DH0101", paragraph_text)
            self.assertGreaterEqual(len(document.tables), 6)

            detail_tables = self._find_tables_with_label(document, "หมวดตัวชี้วัด")
            detail_table = detail_tables[-1]
            detail_labels = [detail_table.cell(row_index, 0).text.strip() for row_index in range(len(detail_table.rows))]
            detail_values = [detail_table.cell(row_index, 1).text.strip() for row_index in range(len(detail_table.rows))]

            self.assertIn("หมวดตัวชี้วัด", detail_labels)
            self.assertIn("ประเภทตัวชี้วัด", detail_labels)
            self.assertIn("รหัสตัวชี้วัด", detail_labels)
            self.assertIn("ชื่อตัวชี้วัด (ภาษาไทย)", detail_labels)
            self.assertIn("ข้อมูลที่ต้องการ", detail_labels)
            self.assertIn("รหัสโรค/หัตถการที่เกี่ยวข้อง", detail_labels)
            self.assertIn("Data Completeness", detail_labels)
            self.assertIn("Validation Summary", detail_labels)
            self.assertIn("PS0301", detail_values)
            all_table_values = "\n".join(
                table.cell(row_index, 1).text.strip()
                for table in document.tables[-3:]
                for row_index in range(len(table.rows))
            )
            self.assertIn("Complete", all_table_values)
            self.assertIn("ข้อมูล required ครบถ้วน", all_table_values)
            self.assertNotIn("Warnings:", all_table_values)
            self.assertNotIn("Data Completeness: Complete", paragraph_text)

            summary_table = document.tables[-1]
            summary_headers = [summary_table.cell(0, col_index).text.strip() for col_index in range(len(summary_table.columns))]
            summary_values = "\n".join(
                summary_table.cell(row_index, col_index).text.strip()
                for row_index in range(len(summary_table.rows))
                for col_index in range(len(summary_table.columns))
            )
            self.assertEqual(summary_headers, ["KPI Code", "KPI Name TH", "Data Completeness", "Validation Summary"])
            self.assertIn("DH0101", summary_values)
            self.assertIn("PS0301", summary_values)

    def test_generate_kpi_dictionary_shows_incomplete_status_and_missing_fields_for_each_kpi(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            template_path = Path(tmpdir) / "template.xlsx"
            output_path = Path(tmpdir) / "KPI_Dictionary.docx"
            report_path = Path(tmpdir) / "validation_report.xlsx"

            create_template_workbook(template_path)
            workbook = load_workbook(template_path)
            input_sheet = workbook[SHEET_NAMES["kpi_input_form"]]
            input_sheet["E2"] = ""
            input_sheet["G2"] = ""
            input_sheet["F2"] = ""
            input_sheet["K2"] = ""
            input_sheet["L2"] = ""
            workbook.save(template_path)

            generate_kpi_dictionary(template_path, output_path, validation_report_path=report_path)

            document = Document(output_path)
            first_detail_table = self._find_tables_with_label(document, "หมวดตัวชี้วัด")[0]
            first_labels = [first_detail_table.cell(row_index, 0).text.strip() for row_index in range(len(first_detail_table.rows))]
            first_values = [first_detail_table.cell(row_index, 1).text.strip() for row_index in range(len(first_detail_table.rows))]

            self.assertIn("Data Completeness", first_labels)
            self.assertIn("Incomplete", "\n".join(first_values))
            self.assertIn("Validation Summary", first_labels)
            self.assertIn("ขาดข้อมูล: นิยาม คำอธิบาย ความหมายของตัวชี้วัด", "\n".join(first_values))
            self.assertNotIn("KPI_Name_EN", "\n".join(first_values))
            self.assertNotIn("Objective_TH", "\n".join(first_values))
            self.assertNotIn("Numerator_CodeSet", "\n".join(first_values))

            summary_table = document.tables[-1]
            summary_values = "\n".join(
                summary_table.cell(row_index, col_index).text.strip()
                for row_index in range(len(summary_table.rows))
                for col_index in range(len(summary_table.columns))
            )
            self.assertIn("Incomplete", summary_values)
            self.assertIn("ขาดข้อมูล: นิยาม คำอธิบาย ความหมายของตัวชี้วัด", summary_values)


if __name__ == "__main__":
    unittest.main()
